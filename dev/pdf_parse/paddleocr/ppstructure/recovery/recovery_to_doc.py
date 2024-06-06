# Copyright (c) 2020 PaddlePaddle Authors. All Rights Reserved.
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import os
import re
import pathlib
from copy import deepcopy
import urllib.parse

from docx import Document
from docx import shared
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

from ppstructure.recovery.table_process import HtmlToDocx

from ppocr.utils.logging import get_logger

logger = get_logger()


def identifyHeaders(layout_boxes):
    header_prefix = {} # 标题高度差：header_tag(#) prefix
    header_res = set()
    
    for region in layout_boxes:
        if  region["type"].lower() in ["title", "header"]:
            title_height = region["bbox"][3] - region["bbox"][1]  # TODO：一个标题如果多行会不准确
            header_res.add(title_height)
            
    sorted_heights = sorted(header_res, reverse=True)
    
    for i, height in enumerate(sorted_heights):
        header_prefix[height] = '#' * (i+1) + ' '
        
    return header_prefix
    
    
def get_region_text(region):
    """获取每一个 region 的所有文本"""
    res = ''
    for line in region["res"]:
        res += line["text"]
    return res


def save_markdown(md_text):
    pathlib.Path("./output.md").write_bytes(md_text.encode())
    

def convert_info_markdown(layout_boxes, save_folder, img_name):
    # 1. 识别标题信息
    header_prefix = identifyHeaders(layout_boxes)
    
    # process each region one by one
    res = ''
    for i, region in enumerate(layout_boxes):
        img_idx = region["img_idx"]
        
        if region["type"].lower() == "figure":
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(
                excel_save_folder, "{}_{}.jpg".format(region["bbox"], img_idx)
            )
            encoded_file_name = urllib.parse.quote(img_path)
            if not os.path.exists(img_path):
                print(f"The file {img_path} does not exists.")

            figure_caption = 'Figure'
            res += "\n"
            if i > 0 and layout_boxes[i-1]["type"].lower() == "figure_caption":
                figure_caption = get_region_text(layout_boxes[i-1])
                res += f"{figure_caption}\n" # 插入描述
                res += f"![{figure_caption}]({str(encoded_file_name)})\n" # 插入图片
            elif i + 1 < len(layout_boxes) and layout_boxes[i+1]["type"].lower() == "figure_caption":
                figure_caption = get_region_text(layout_boxes[i+1])
                res += f"![{figure_caption}]({str(encoded_file_name)})\n" # 插入图片
                res += f"{figure_caption}\n" # 插入描述
            else:
                res += f"![{figure_caption}]({str(encoded_file_name)})\n"
            
        elif region["type"].lower() in ["title", "header"]:
            title_height = region["bbox"][3] - region["bbox"][1]
            hdr_string = header_prefix.get(title_height, '')
            res += hdr_string + get_region_text(region) + "\n\n"
    
        elif region["type"].lower() == "table": # TODO: customize style for table
            table_content = re.search(r'(<table.*?</table>)', region["res"]["html"], re.DOTALL).group(1)
            res += f"{table_content}\n\n"
            
        elif region["type"].lower() == "table_caption":
            table_caption = get_region_text(region)
            res += f"{table_caption}\n"
        
        elif region["type"].lower() == "text":
            res += get_region_text(region) + "\n"
            
    save_markdown(res)
    return res

def convert_info_docx(img, res, save_folder, img_name):
    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    doc.styles["Normal"].font.size = shared.Pt(6.5)

    flag = 1 # 当前分栏状态：1为单栏；2为双栏。
    for i, region in enumerate(res):
        if len(region["res"]) == 0 and region["type"].lower() != "figure":
            continue
        img_idx = region["img_idx"]
        if flag == 2 and region["layout"] == "single":
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "1")
            flag = 1
        elif flag == 1 and region["layout"] == "double":
            section = doc.add_section(WD_SECTION.CONTINUOUS)
            section._sectPr.xpath("./w:cols")[0].set(qn("w:num"), "2")
            flag = 2

        if region["type"].lower() == "figure":
            excel_save_folder = os.path.join(save_folder, img_name)
            img_path = os.path.join(
                excel_save_folder, "{}_{}.jpg".format(region["bbox"], img_idx)
            )
            if os.path.exists(img_path):
                print(f"The file {img_path} exists.")
            else:
                print(f"The file {img_path} does not exist.")
            paragraph_pic = doc.add_paragraph()
            paragraph_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph_pic.add_run("")
            if flag == 1:
                run.add_picture(img_path, width=shared.Inches(5))
            elif flag == 2:
                run.add_picture(img_path, width=shared.Inches(2))
        elif region["type"].lower() == "title":
            doc.add_heading(region["res"][0]["text"])
        elif region["type"].lower() == "table":
            parser = HtmlToDocx()
            parser.table_style = "TableGrid"
            parser.handle_table(region["res"]["html"], doc)
        else:
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            for i, line in enumerate(region["res"]):
                if i == 0:
                    paragraph_format.first_line_indent = shared.Inches(0.25)
                text_run = paragraph.add_run(line["text"] + " ")
                text_run.font.size = shared.Pt(10)

    # save to docx
    docx_path = os.path.join(save_folder, "{}_ocr.docx".format(img_name))
    doc.save(docx_path)
    logger.info("docx save to {}".format(docx_path))


def sorted_layout_boxes(res, w):
    """
    Sort text boxes in order from top to bottom, left to right
    args:
        res(list):ppstructure results
    return:
        sorted results(list)
    """
    num_boxes = len(res)
    if num_boxes == 1:
        res[0]["layout"] = "single"
        return res

    sorted_boxes = sorted(res, key=lambda x: (x["bbox"][1], x["bbox"][0]))
    _boxes = list(sorted_boxes)

    new_res = []
    res_left = []
    res_right = []
    i = 0

    while True:
        if i >= num_boxes:
            break
        if i == num_boxes - 1:
            if (
                _boxes[i]["bbox"][1] > _boxes[i - 1]["bbox"][3]
                and _boxes[i]["bbox"][0] < w / 2
                and _boxes[i]["bbox"][2] > w / 2
            ):
                new_res += res_left
                new_res += res_right
                _boxes[i]["layout"] = "single"
                new_res.append(_boxes[i])
            else:
                if _boxes[i]["bbox"][2] > w / 2:
                    _boxes[i]["layout"] = "double"
                    res_right.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
                elif _boxes[i]["bbox"][0] < w / 2:
                    _boxes[i]["layout"] = "double"
                    res_left.append(_boxes[i])
                    new_res += res_left
                    new_res += res_right
            res_left = []
            res_right = []
            break
        elif _boxes[i]["bbox"][2] <= w / 2:
            _boxes[i]["layout"] = "double"
            res_left.append(_boxes[i])
            i += 1
        elif _boxes[i]["bbox"][0] >= w / 2:
            _boxes[i]["layout"] = "double"
            res_right.append(_boxes[i])
            i += 1
        else: #  _boxes[i]["bbox"][0] < w / 2 and _boxes[i]["bbox"][0] > w / 2
            new_res += res_left
            new_res += res_right
            _boxes[i]["layout"] = "single"
            new_res.append(_boxes[i])
            res_left = []
            res_right = []
            i += 1
    if res_left:
        new_res += res_left
    if res_right:
        new_res += res_right
    return new_res
